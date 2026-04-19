"""
Run this script to generate the project report Word document.
    python generate_report.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ────────────────────────────────────────────────────────────
PRIMARY   = RGBColor(0x1A, 0x73, 0xE8)   # Google Blue
ACCENT    = RGBColor(0x0D, 0x47, 0xA1)   # Dark Blue
LIGHT_BG  = RGBColor(0xF0, 0xF4, 0xFF)   # Light blue tint
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT = RGBColor(0x1A, 0x1A, 0x2E)
MID_GRAY  = RGBColor(0x55, 0x65, 0x7A)
GREEN     = RGBColor(0x1E, 0x88, 0x55)
ORANGE    = RGBColor(0xE6, 0x51, 0x00)


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color", "1A73E8"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_para(doc, text="", bold=False, italic=False, size=11,
             color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.color.rgb = color or DARK_TEXT
    return p


def add_heading(doc, text, level=1):
    sizes   = {1: 22, 2: 16, 3: 13}
    colors  = {1: PRIMARY, 2: ACCENT, 3: DARK_TEXT}
    bolds   = {1: True, 2: True, 3: True}
    before  = {1: 18, 2: 14, 3: 10}
    after   = {1: 6, 2: 4, 3: 3}

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before[level])
    p.paragraph_format.space_after  = Pt(after[level])
    run = p.add_run(text)
    run.bold = bolds[level]
    run.font.size  = Pt(sizes[level])
    run.font.color.rgb = colors[level]

    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "1A73E8")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = color or DARK_TEXT
    return p


def add_info_box(doc, label, content, label_color="1A73E8"):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Inches(1.5)
    tbl.columns[1].width = Inches(5.0)

    lc = tbl.rows[0].cells[0]
    rc = tbl.rows[0].cells[1]
    set_cell_bg(lc, label_color)
    set_cell_bg(rc, "F0F4FF")

    lp = lc.paragraphs[0]
    lr = lp.add_run(label)
    lr.bold = True
    lr.font.size = Pt(10)
    lr.font.color.rgb = WHITE
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    rp = rc.paragraphs[0]
    rr = rp.add_run(content)
    rr.font.size = Pt(10)
    rr.font.color.rgb = DARK_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def add_two_col_table(doc, headers, rows, header_bg="1A73E8"):
    num_cols = len(headers)
    num_rows = 1 + len(rows)
    tbl = doc.add_table(rows=num_rows, cols=num_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            if ci >= num_cols:
                break
            cell = tbl.rows[ri + 1].cells[ci]
            bg = "F8FAFF" if ri % 2 == 0 else "FFFFFF"
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.color.rgb = DARK_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl


def build_report():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ══════════════════════════════════════════════════════════════════════════
    #  COVER PAGE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph().paragraph_format.space_after = Pt(40)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t1 = title_p.add_run("Intelligent Shopping Advisor")
    t1.bold = True
    t1.font.size = Pt(32)
    t1.font.color.rgb = PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s1 = sub_p.add_run("AI-Powered Multi-Agent Product Recommendation System")
    s1.font.size = Pt(16)
    s1.italic = True
    s1.font.color.rgb = MID_GRAY

    doc.add_paragraph().paragraph_format.space_after = Pt(30)

    # Divider line
    div = doc.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div.paragraph_format.space_after = Pt(30)
    divrun = div.add_run("─" * 60)
    divrun.font.color.rgb = PRIMARY

    # Cover info table
    cover_tbl = doc.add_table(rows=6, cols=2)
    cover_tbl.style = "Table Grid"
    cover_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    cover_data = [
        ("Student Name",   "Muhammad Abdul Wasay"),
        ("Student ID",     "L1F22BSAI0006"),
        ("Subject",        "Generative Artificial Intelligence"),
        ("Assignment",     "Assignment 2"),
        ("Submitted To",   "Course Instructor"),
        ("Date",           "April 2026"),
    ]
    for ri, (label, value) in enumerate(cover_data):
        lc = cover_tbl.rows[ri].cells[0]
        vc = cover_tbl.rows[ri].cells[1]
        set_cell_bg(lc, "1A73E8")
        set_cell_bg(vc, "F0F4FF" if ri % 2 == 0 else "FFFFFF")

        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(11)
        lr.font.color.rgb = WHITE
        lc.width = Inches(2.2)

        vp = vc.paragraphs[0]
        vr = vp.add_run(value)
        vr.font.size = Pt(11)
        vr.font.color.rgb = DARK_TEXT

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #  1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. Executive Summary")
    add_para(doc, (
        "The Intelligent Shopping Advisor is a production-grade, AI-powered multi-agent system "
        "designed to assist users in making informed product purchase decisions. Built using "
        "LangGraph for multi-agent orchestration, Google Gemini 2.5 Flash as the reasoning LLM, "
        "the Rainforest API for live Amazon product retrieval, and MongoDB Atlas for intelligent "
        "caching, the system provides end-to-end shopping assistance through a conversational "
        "Chainlit interface."
    ), size=10.5, space_after=8)
    add_para(doc, (
        "A user simply types a natural language query such as 'Best gaming laptop under $1000' "
        "and the system orchestrates four specialized AI agents in sequence — extracting "
        "preferences, fetching live products from Amazon, scoring them against user requirements, "
        "and generating personalized recommendations complete with product images, prices, ratings, "
        "and justifications."
    ), size=10.5, space_after=8)

    # ══════════════════════════════════════════════════════════════════════════
    #  2. PROJECT OVERVIEW
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. Project Overview & Objectives")
    add_para(doc, (
        "Modern e-commerce platforms offer millions of products, making it difficult for consumers "
        "to identify the best option for their specific needs and budget. This project addresses "
        "this challenge by building an AI advisor that simulates expert-level product consultation "
        "through the following capabilities:"
    ), size=10.5, space_after=6)

    objectives = [
        "Accept natural language shopping queries with no structured input required",
        "Automatically extract budget, product category, and feature priorities using LLM reasoning",
        "Retrieve live, real-time Amazon product data through the Rainforest API",
        "Compare products across price fit, feature relevance, and customer ratings",
        "Generate personalized, justified recommendations with product images and buy links",
        "Cache retrieved products in MongoDB Atlas for efficient repeat query handling",
        "Support follow-up questions and multi-turn conversation through session memory",
    ]
    for obj in objectives:
        add_bullet(doc, obj)

    # ══════════════════════════════════════════════════════════════════════════
    #  3. SYSTEM ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. System Architecture")

    add_heading(doc, "3.1 High-Level Architecture", level=2)
    add_para(doc, (
        "The system follows a pipeline architecture where each user query flows through four "
        "AI agents in sequence, managed by a LangGraph StateGraph. The shared state object "
        "carries data between agents, enabling each agent to build on the previous one's output."
    ), size=10.5, space_after=8)

    arch_data = [
        ("User Interface",     "Chainlit Chat Application",      "Python · Chainlit 2.x"),
        ("Orchestration",      "LangGraph State Machine",         "LangGraph · LangChain"),
        ("AI Reasoning",       "Google Gemini LLM",               "Gemini 2.5 Flash"),
        ("Product Data",       "Live Amazon Search",              "Rainforest API"),
        ("Data Persistence",   "Product Cache + History",         "MongoDB Atlas"),
        ("Semantic Search",    "Vector Similarity Index",         "FAISS · Sentence-Transformers"),
    ]
    add_two_col_table(doc,
        ["Layer", "Component", "Technology"],
        arch_data
    )

    add_heading(doc, "3.2 LangGraph Workflow", level=2)
    add_para(doc, (
        "The LangGraph StateGraph defines the pipeline as a directed graph with nodes "
        "(agents) and edges (transitions). A conditional edge after the Retrieval Agent "
        "routes to an error handler if no products are found, otherwise continues to "
        "Comparison."
    ), size=10.5, space_after=6)

    flow_items = [
        ("START", "User query received", "0D47A1"),
        ("Node 1", "Preference Agent  →  Extracts structured preferences from query", "1A73E8"),
        ("Node 2", "Retrieval Agent   →  Fetches live Amazon products (Rainforest API)", "1565C0"),
        ("Condition", "Products found?  →  YES: Comparison   |   NO: Error Handler", "E65100"),
        ("Node 3", "Comparison Agent  →  Scores products (0-10 across 3 dimensions)", "1A73E8"),
        ("Node 4", "Recommendation Agent  →  Generates personalized write-ups", "1565C0"),
        ("END", "Results displayed in Chainlit with images and buy links", "1E8855"),
    ]
    for stage, desc, color in flow_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Inches(0.3)
        r1 = p.add_run(f"{stage}: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor(
            int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        )
        r2 = p.add_run(desc)
        r2.font.size = Pt(10)
        r2.font.color.rgb = DARK_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ══════════════════════════════════════════════════════════════════════════
    #  4. MULTI-AGENT DESIGN
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "4. Multi-Agent System Design")

    agents = [
        (
            "Agent 1 — Preference Agent",
            "agents/preference_agent.py",
            (
                "The Preference Agent is the first node in the pipeline. It receives the raw "
                "natural language query from the user and uses Google Gemini 2.5 Flash to perform "
                "structured information extraction. The agent outputs a JSON object containing the "
                "product category, budget range (in USD), must-have features, nice-to-have features, "
                "brand preferences, and optimized Amazon search terms."
            ),
            [
                "Input: Raw user query string + conversation history",
                "LLM: Gemini 2.5 Flash with temperature=0 for deterministic extraction",
                "Currency conversion: PKR to USD (÷280) handled automatically",
                "Budget inference: 'cheap' → $200, 'mid-range' → $600, 'premium' → $800+",
                "Fallback: Basic keyword extraction if Gemini is unavailable",
                "Output: Structured preferences dict added to shared state",
            ]
        ),
        (
            "Agent 2 — Retrieval Agent",
            "agents/retrieval_agent.py",
            (
                "The Retrieval Agent fetches live product data from Amazon using the Rainforest "
                "API. It uses the search terms generated by the Preference Agent to make real-time "
                "API calls, returning up to 20 products per search term. Results are cached in "
                "MongoDB Atlas to avoid redundant API calls on repeat queries. Price and brand "
                "filters are applied post-retrieval."
            ),
            [
                "Input: Structured preferences (category, budget, brand, search terms)",
                "Primary source: Rainforest API (live Amazon.com search results)",
                "Caching: Results stored in MongoDB Atlas with unique ASIN index",
                "Filters: Price range and brand applied to API results",
                "Fallback: MongoDB cache queried if API returns empty results",
                "Output: Up to 20 candidate products with image URLs and Amazon links",
            ]
        ),
        (
            "Agent 3 — Comparison Agent",
            "agents/comparison_agent.py",
            (
                "The Comparison Agent evaluates each candidate product against the user's "
                "preferences using Gemini. It assigns three dimension scores per product and "
                "computes a weighted overall score. The scoring prompt is carefully engineered "
                "to produce consistent, calibrated 0-10 scores. Products are returned ranked "
                "by their overall score."
            ),
            [
                "Input: Candidate products + user preferences",
                "LLM: Gemini 2.5 Flash with temperature=0",
                "Scoring dimensions: Price fit (30%), Feature match (45%), Rating quality (25%)",
                "Handles missing prices gracefully (neutral score of 5)",
                "Fallback: Rating-based sort if Gemini call fails",
                "Output: Scored and ranked product list with match_reason per product",
            ]
        ),
        (
            "Agent 4 — Recommendation Agent",
            "agents/recommendation_agent.py",
            (
                "The Recommendation Agent takes the top 5 ranked products and generates "
                "personalized, human-readable justifications for each. It uses Gemini with "
                "slightly higher temperature (0.3) to produce natural, varied language. The "
                "output is a structured JSON array with per-product fields for why_its_right, "
                "best_for, and trade_off — displayed as individual product cards in the UI."
            ),
            [
                "Input: Top 5 ranked products + original user query + preferences",
                "LLM: Gemini 2.5 Flash with temperature=0.3 for natural language",
                "Output format: Structured JSON array (not freeform text)",
                "Each product gets: why_its_right, best_for, trade_off fields",
                "Fallback: Basic price/rating summary if generation fails",
                "UI: Cards streamed word-by-word with inline product image",
            ]
        ),
    ]

    for agent_name, file_path, description, bullets in agents:
        add_heading(doc, agent_name, level=2)
        add_para(doc, f"File: {file_path}", italic=True, size=9.5, color=MID_GRAY, space_after=4)
        add_para(doc, description, size=10.5, space_after=6)
        for b in bullets:
            add_bullet(doc, b)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ══════════════════════════════════════════════════════════════════════════
    #  5. DATA HANDLING
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "5. Data Handling & Storage")

    add_heading(doc, "5.1 Rainforest API — Live Product Data", level=2)
    add_para(doc, (
        "The Rainforest API provides structured access to Amazon product data without "
        "requiring direct web scraping. Each query triggers a live search on Amazon.com, "
        "returning up to 20 products with full metadata."
    ), size=10.5, space_after=6)

    add_two_col_table(doc,
        ["Field Returned", "Usage in System"],
        [
            ("ASIN (Product ID)",       "Unique identifier for MongoDB deduplication"),
            ("Title",                   "Product name displayed in recommendation card"),
            ("Brand",                   "Brand preference filtering + display"),
            ("Price (value + currency)", "Budget filtering + price display"),
            ("Rating (0–5 stars)",      "Rating score calculation + display"),
            ("Ratings Total",           "Social proof shown in card"),
            ("Image URL",               "Inline product image in Chainlit UI"),
            ("Product Link",            "View on Amazon clickable button"),
        ]
    )

    add_heading(doc, "5.2 MongoDB Atlas — Caching Layer", level=2)
    add_para(doc, (
        "MongoDB Atlas (free M0 tier) serves as an intelligent cache. Every product "
        "fetched from the Rainforest API is stored with a unique index on the ASIN field "
        "to prevent duplicates. This reduces API credit usage on repeat queries and "
        "improves response time for popular product categories."
    ), size=10.5, space_after=6)

    add_para(doc, "Product document schema:", bold=True, size=10, space_after=3)
    schema_fields = [
        ("asin",          "String",   "Unique Amazon product identifier"),
        ("name",          "String",   "Full product title"),
        ("brand",         "String",   "Manufacturer brand"),
        ("price",         "Float",    "Price in USD"),
        ("currency",      "String",   "Currency code (USD)"),
        ("rating",        "Float",    "Average star rating (0.0–5.0)"),
        ("ratings_total", "Integer",  "Total number of customer reviews"),
        ("image_url",     "String",   "CDN URL of product image"),
        ("product_url",   "String",   "Full Amazon product page URL"),
        ("category",      "String",   "Product category assigned by system"),
        ("features",      "Array",    "List of key product features"),
        ("fetched_at",    "DateTime", "ISO timestamp of when fetched"),
    ]
    add_two_col_table(doc,
        ["Field", "Type", "Description"],
        schema_fields
    )

    add_heading(doc, "5.3 FAISS Vector Index — Semantic Search", level=2)
    add_para(doc, (
        "A FAISS (Facebook AI Similarity Search) index is built over cached MongoDB products "
        "using the all-MiniLM-L6-v2 sentence-transformer model. This enables semantic "
        "matching — a query like 'phone for photography' retrieves products with camera-related "
        "features even when exact keywords don't match. The index is saved to disk and "
        "reloaded on startup."
    ), size=10.5, space_after=6)

    # ══════════════════════════════════════════════════════════════════════════
    #  6. USER INTERFACE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "6. User Interface — Chainlit Chat Application")
    add_para(doc, (
        "The frontend is built with Chainlit, an open-source Python framework for "
        "building production-grade conversational AI interfaces. The UI closely resembles "
        "ChatGPT with additional rich media support."
    ), size=10.5, space_after=6)

    ui_features = [
        "Natural language chat interface — no forms or dropdowns required",
        "Collapsible agent step panels showing each agent's input and output in real time",
        "Per-agent status messages (Agent 1 done, Agent 2 done...) for full transparency",
        "Product cards streamed word-by-word in typewriter style for a modern feel",
        "Inline product images displayed inside chat messages using cl.Image",
        "Clickable View on Amazon links for each recommended product",
        "Multi-turn conversation memory — follow-up questions retain context",
        "Session-based user state management via cl.user_session",
    ]
    for f in ui_features:
        add_bullet(doc, f)

    # ══════════════════════════════════════════════════════════════════════════
    #  7. REASONING & EXPLAINABILITY
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "7. Reasoning & Explainability")
    add_para(doc, (
        "A key strength of this system is its ability to explain every decision at each "
        "stage of the pipeline. Unlike black-box recommendation engines, this system "
        "provides full transparency through:"
    ), size=10.5, space_after=6)

    explain_items = [
        ("Preference Extraction",
         "The system shows exactly what it understood from your query — category, budget, and features — before searching."),
        ("Retrieval Transparency",
         "The number of live products fetched is shown, along with the search terms used on Amazon."),
        ("Scoring Breakdown",
         "Each product receives a visible score breakdown: price fit, feature match, and rating quality with weights."),
        ("Justification per Product",
         "Gemini writes a specific 'Why it's right for you' paragraph per product referencing your stated budget and features."),
        ("Trade-off Honesty",
         "Each recommendation includes an honest Trade-off so users understand the product's limitations."),
        ("Best For",
         "A 'Best for' label identifies the ideal buyer profile, helping users self-select the right option."),
    ]

    for label, text in explain_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Inches(0.3)
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = PRIMARY
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = DARK_TEXT

    # ══════════════════════════════════════════════════════════════════════════
    #  8. TESTING & EVALUATION
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "8. Testing & Evaluation")

    add_heading(doc, "8.1 Test Dataset", level=2)
    add_para(doc, (
        "A comprehensive test dataset of 55 queries was created covering six categories "
        "of test scenarios. The dataset is stored in tests/test_queries.json and used "
        "by the automated evaluation script tests/evaluate.py."
    ), size=10.5, space_after=6)

    add_two_col_table(doc,
        ["Query Category", "Count", "Description", "Example"],
        [
            ("Normal",        "20", "Typical well-formed queries",        "Best smartphone under $500"),
            ("PKR Budget",    "5",  "Pakistani Rupee budget queries",     "Laptop under 150,000 PKR"),
            ("Vague",         "5",  "Minimal context queries",            "I need a new phone"),
            ("Multi-Criteria","5",  "Complex multi-requirement queries",  "Gaming laptop under $1200 with RTX GPU"),
            ("Edge Cases",    "10", "Boundary and invalid inputs",        "Laptop under $1 / conflicting specs"),
            ("Follow-ups",    "3",  "Conversational follow-up queries",   "What about a cheaper option?"),
            ("Comparisons",   "2",  "Direct product comparison queries",  "iPhone 15 vs Samsung S24"),
            ("Extra Normal",  "5",  "Additional standard queries",        "Camera for portrait photography"),
        ]
    )

    add_heading(doc, "8.2 Evaluation Metrics", level=2)
    add_two_col_table(doc,
        ["Metric", "Description", "Measurement Method"],
        [
            ("Category Accuracy",  "Correct product category extracted",     "Expected vs actual category field"),
            ("Budget Accuracy",    "Budget within 15% of stated value",      "Numeric comparison with tolerance"),
            ("Recommendation Rate","Queries that returned recommendations",   "Boolean: recommendations list non-empty"),
            ("Image Availability", "Products returned with valid image URLs", "URL presence check per product"),
            ("Response Time",      "End-to-end pipeline execution time",      "Python time.time() measurement (seconds)"),
            ("Relevance Score",    "AI-assigned match score",                 "Comparison agent overall_score average"),
        ]
    )

    add_heading(doc, "8.3 Running the Evaluation", level=2)
    code_paras = [
        "# Run all 55 test queries",
        "python -m tests.evaluate",
        "",
        "# Run only normal queries",
        "python -m tests.evaluate --category normal",
        "",
        "# Run first 10 queries only",
        "python -m tests.evaluate --limit 10",
    ]
    for line in code_paras:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Inches(0.4)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ══════════════════════════════════════════════════════════════════════════
    #  9. PROJECT STRUCTURE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "9. Project File Structure")

    file_structure = [
        ("app.py",                        "Chainlit application — UI, agent orchestration, streaming output"),
        ("config.py",                     "Central configuration — loads .env, defines constants"),
        ("requirements.txt",              "Python package dependencies with version constraints"),
        (".env",                          "API keys and MongoDB URI (not committed to version control)"),
        ("rainforest/client.py",          "Rainforest API wrapper — search_products() function"),
        ("database/mongo_client.py",      "MongoDB Atlas connection, CRUD operations, filtering"),
        ("database/seed.py",              "Optional cache pre-warmer — seeds product categories"),
        ("vector_store/faiss_index.py",   "FAISS semantic index — build, save, load, search"),
        ("agents/state.py",               "LangGraph ShoppingState TypedDict definition"),
        ("agents/preference_agent.py",    "Agent 1 — Gemini-based preference extraction"),
        ("agents/retrieval_agent.py",     "Agent 2 — Live Rainforest API + MongoDB cache"),
        ("agents/comparison_agent.py",    "Agent 3 — Gemini product scoring (3 dimensions)"),
        ("agents/recommendation_agent.py","Agent 4 — Gemini personalized justification generation"),
        ("graph/workflow.py",             "LangGraph StateGraph — nodes, edges, conditional routing"),
        ("tests/test_queries.json",       "55 test queries across 8 categories"),
        ("tests/evaluate.py",             "Automated evaluation — accuracy, relevance, response time"),
    ]

    add_two_col_table(doc,
        ["File / Module", "Purpose"],
        file_structure
    )

    # ══════════════════════════════════════════════════════════════════════════
    #  10. HOW TO RUN
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "10. Setup & Running the Application")

    steps = [
        ("Step 1 — Prerequisites",
         "Python 3.10+, pip, MongoDB Atlas account (free M0 tier), "
         "Google AI Studio API key, Rainforest API key."),
        ("Step 2 — Install Dependencies",
         "pip install -r requirements.txt"),
        ("Step 3 — Configure Environment",
         "Fill in .env with: RAINFOREST_API_KEY, GEMINI_API_KEY, MONGODB_URI, MONGODB_DB_NAME"),
        ("Step 4 — Launch Application",
         "python -m chainlit run app.py"),
        ("Step 5 — Open Browser",
         "Navigate to http://localhost:8000 and start chatting"),
    ]

    for step_label, step_desc in steps:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r1 = p.add_run(f"{step_label}: ")
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = PRIMARY
        r2 = p.add_run(step_desc)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = DARK_TEXT

    # ══════════════════════════════════════════════════════════════════════════
    #  11. ETHICAL CONSIDERATIONS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "11. Ethical & Practical Considerations")

    ethical_items = [
        ("Data Privacy",
         "No personal user data is stored. MongoDB only stores product information from "
         "public Amazon listings. Conversation history is session-only and cleared when "
         "the browser tab is closed."),
        ("API Compliance",
         "The Rainforest API is used as an authorized intermediary for Amazon data access, "
         "fully compliant with Amazon's terms of service. No direct scraping of Amazon "
         "pages is performed."),
        ("Bias in Recommendations",
         "Product scoring is based on price fit, feature match, and customer ratings. "
         "Products with no pricing information receive a neutral score (5/10) to avoid "
         "penalizing them unfairly. Brand bias is only introduced when the user explicitly "
         "states a brand preference."),
        ("LLM Limitations",
         "Gemini may occasionally generate inaccurate justifications if product names are "
         "ambiguous. The system mitigates this by instructing the model never to invent "
         "specifications not visible in the product name, and by using temperature=0 for "
         "factual extraction tasks."),
        ("Outdated Product Data",
         "Prices and availability on Amazon change frequently. The system displays a "
         'fetched_at timestamp and links directly to Amazon so users can verify current '
         "pricing before purchase. Cached products older than 24 hours should be refreshed."),
        ("Recommendation Transparency",
         "The system makes its reasoning explicit at every step — users can see exactly "
         "what preferences were extracted, how many products were found, and why each "
         "recommendation was selected. This prevents the black-box effect common in "
         "traditional recommendation systems."),
    ]

    for label, text in ethical_items:
        add_heading(doc, label, level=3)
        add_para(doc, text, size=10.5, space_after=6)

    # ══════════════════════════════════════════════════════════════════════════
    #  12. CONCLUSION
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "12. Conclusion")
    add_para(doc, (
        "The Intelligent Shopping Advisor successfully demonstrates the power of multi-agent "
        "AI systems for real-world decision support. By combining LangGraph's structured "
        "orchestration, Google Gemini's language understanding, live Amazon data through "
        "the Rainforest API, and MongoDB's caching layer, the system delivers fast, "
        "accurate, and explainable product recommendations through an intuitive chat interface."
    ), size=10.5, space_after=8)
    add_para(doc, (
        "The modular agent design ensures each component can be improved or replaced "
        "independently — for example, swapping in a different LLM, adding a new data "
        "source, or extending the comparison logic with additional scoring dimensions. "
        "The system serves as a strong foundation for a production shopping assistant "
        "and demonstrates practical mastery of generative AI system design."
    ), size=10.5, space_after=8)

    # ── Save ──────────────────────────────────────────────────────────────────
    output_path = "Intelligent_Shopping_Advisor_Report.docx"
    doc.save(output_path)
    print(f"\nReport saved: {output_path}\n")


if __name__ == "__main__":
    build_report()
